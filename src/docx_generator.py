"""Render cover letters from DOCX templates using placeholder replacement.

This module takes a generated cover letter body and a user profile, then
fills a German ``.docx`` template by replacing ``[placeholder]`` strings with
actual content. It handles the common DOCX issue where text is split across
multiple XML ``<w:r>`` runs.

The main entry point is :func:`render_cover_letter`, which relies on the
helper functions :func:`_replace_text_in_paragraph` and
:func:`_replace_text_in_docx` to perform the actual run-aware text
replacement throughout the document's paragraphs and tables.
"""

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph

from src.models import JobDescription
from src.utils import get_todays_date


_GENERIC_GREETINGS = {
    "en": "Dear hiring manager,",
    "de": "Sehr geehrte Damen und Herren,",
    "es": "A quien corresponda,",
}


def _replace_text_in_paragraph(paragraph: Paragraph, old_text: str, new_text: str) -> None:
    """Replace ``old_text`` with ``new_text`` in a single paragraph.

    Handles the common DOCX issue where visible text is split across
    multiple ``<w:r>`` run elements by concatenating all runs first,
    performing the replacement, then writing the result back into the
    first run.

    Args:
        paragraph: The python-docx Paragraph object to modify in place.
        old_text: The placeholder string to search for (e.g. ``"[date]"``).
        new_text: The actual content to replace the placeholder with.
    """
    full_text = "".join(run.text for run in paragraph.runs)

    if old_text not in full_text:
        return

    replaced = full_text.replace(old_text, new_text)

    if not paragraph.runs:
        return

    for run in paragraph.runs:
        run.text = ""

    paragraph.runs[0].text = replaced


def _replace_text_in_docx(doc: Document, old_text: str, new_text: str) -> None:
    """Replace ``old_text`` with ``new_text`` throughout an entire document.

    Iterates over every paragraph in the document body and inside all
    tables, delegating to :func:`_replace_text_in_paragraph` for each one.

    Args:
        doc: The python-docx Document object to modify in place.
        old_text: The placeholder string to search for.
        new_text: The actual content to replace the placeholder with.
    """
    for paragraph in doc.paragraphs:
        _replace_text_in_paragraph(paragraph, old_text, new_text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, old_text, new_text)


def _build_greeting(job: JobDescription, language: str) -> str:
    """Build the salutation line for a cover letter.

    Uses a personalised formal greeting when *job.receiver_name* is present,
    otherwise falls back to a language-specific generic greeting.

    For German and Spanish, the function parses gender indicators that the LLM
    prepended to the stored name (e.g. ``"Herr Dr. Mustermann"`` or
    ``"Señora López"``) and adapts the adjective accordingly.

    Args:
        job: The job description, optionally including a contact person.
        language: Language code (``"en"``, ``"de"``, ``"es"``).

    Returns:
        The complete greeting line, e.g. ``"Sehr geehrter Herr Dr. Mustermann,"`` or
        ``"Sehr geehrte Damen und Herren,"``.
    """
    if not job.receiver_name:
        return _GENERIC_GREETINGS[language]

    name = job.receiver_name.strip()

    if language == "de":
        if name.startswith("Herr"):
            return f"Sehr geehrter {name},"
        elif name.startswith("Frau"):
            return f"Sehr geehrte {name},"
        else:
            return _GENERIC_GREETINGS["de"]

    if language == "es":
        if name.startswith("Señor "):
            return f"Estimado {name},"
        elif name.startswith("Señora ") or name.startswith("Señorita "):
            return f"Estimada {name},"
        else:
            return _GENERIC_GREETINGS["es"]

    # English and any other language – keep it simple
    return f"Dear {name},"


def _build_cover_letter_context(
    job: JobDescription,
    letter_body: str,
    language: str,
) -> dict[str, str]:
    """Build a context dictionary mapping template placeholders to values.

    Combines data from the job description and generated
    letter body into a flat dictionary whose keys match the ``[placeholder]``
    strings used in the DOCX template. The dictionary includes date, company,
    location, greeting, and letter body.

    Args:
        job: The job description containing company, location, etc.
        letter_body: The AI-generated cover letter body text.
        language: Language code used for date formatting (``"en"``, ``"de"``, ``"es"``).

    Returns:
        A dictionary mapping ``[placeholder]`` keys to their replacement values.
    """
    return {
        "[date]": get_todays_date(language),
        "[company_name]": job.company,
        "[location]": job.location,
        "[greeting]": _build_greeting(job, language),
        "[letter_body]": letter_body,
    }


def render_cover_letter(
    template_path: str | Path,
    output_path: str | Path,
    job: JobDescription,
    letter_body: str,
    language: str,
) -> str:
    """Render a cover letter from a DOCX template by replacing placeholders.

    Loads the template document, replaces all ``[placeholder]`` strings with
    actual values from the job description and user profile, then saves the
    result to the specified output path.

    Args:
        template_path: Filesystem path to the ``.docx`` template file.
        output_path: Filesystem path where the rendered document will be saved.
        job: The job description to pull company and location data from.
        letter_body: The AI-generated cover letter body text.
        language: Language code for date formatting (``"en"``, ``"de"``, ``"es"``).

    Returns:
        The ``output_path`` string where the document was saved.

    Raises:
        FileNotFoundError: If the template file does not exist at
            *template_path*.
        OSError: If the output directory cannot be created or the document
            cannot be saved to *output_path*.
    """
    doc = Document(str(template_path))

    context = _build_cover_letter_context(job, letter_body, language)

    for placeholder, value in context.items():
        _replace_text_in_docx(doc, placeholder, value)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc.save(str(output_path))

    return str(output_path)
